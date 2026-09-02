"""DOCX-Renderer für KI-Briefings und KI-Feedback.

Layout-Basis ist die Briefing-Vorlage der Kursleitung ("BWL A_Briefing
TPn.docx", HSG-Briefvorlage): ``backend/config/ki_rubrics/briefing_template.docx``
ist davon abgeleitet — Stile (Gill Sans MT Pro Light für Überschriften,
Palatino Linotype im Fliesstext, nummerierte Überschriften), Seitenränder und
die Fusszeile "Touchpoint n | BWL A HS2026 | Seite" bleiben, der Inhalt und
das schwere Titelbild wurden entfernt. Fehlt die Vorlage, fällt der Renderer
auf ein neutrales Dokument zurück.

Produkt 1 (Briefing, an die ÜGL): ein Dokument je Übungsgruppe mit einem
nummerierten Abschnitt je Stammgruppe — Kenndaten, formale Vorprüfung
(gemeldet, nicht bewertet), je Baustein Kernposition, tragende Argumente,
dünne Stellen als Rückfrage-Ansatz und die Einschätzung in Prosa. Keine
Punkte, keine Stufen, keine interne Kriterien-Einstufung.

Produkt 2 (Feedback, an die Stammgruppe): ein Dokument je Stammgruppe — je
Baustein was trägt / was bleibt dünn / nächster Schritt, Abschluss Ausblick.
"""

from __future__ import annotations

import io
from datetime import date, datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from backend.briefings.rubrics import KI_RUBRICS_DIR, BRIEFING_SCHEDULE, BriefingRubric

TEMPLATE_PATH: Path = KI_RUBRICS_DIR / "briefing_template.docx"
_GREY = RGBColor(0x59, 0x59, 0x59)

_MONTHS_DE = [
    "Januar", "Februar", "März", "April", "Mai", "Juni",
    "Juli", "August", "September", "Oktober", "November", "Dezember",
]


def _fmt_date(value: date | None) -> str:
    return value.strftime("%d.%m.%Y") if value else "–"


def _fmt_date_long(value: date) -> str:
    return f"{value.day}. {_MONTHS_DE[value.month - 1]} {value.year}"


def _new_document():
    return Document(str(TEMPLATE_PATH)) if TEMPLATE_PATH.exists() else Document()


def _style(doc, *candidates: str) -> str | None:
    names = {s.name for s in doc.styles}
    for name in candidates:
        if name in names:
            return name
    return None


def _para(doc, text: str, *, style: str | None = None, bold_label: str | None = None,
          italic: bool = False, size: float | None = None, grey: bool = False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if bold_label:
        run = p.add_run(bold_label + " ")
        run.bold = True
        if size:
            run.font.size = Pt(size)
        if grey:
            run.font.color.rgb = _GREY
    run = p.add_run(text)
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if grey:
        run.font.color.rgb = _GREY
    return p


def _heading(doc, text: str, level: int) -> None:
    # Vorlage nutzt Heading 2/3 als nummerierte Gliederung (1, 1.1 …)
    doc.add_heading(text, level=level)


def _bullets(doc, items: list[str], empty_text: str) -> None:
    if not items:
        _para(doc, empty_text, italic=True)
        return
    style = _style(doc, "Aufzählung", "List Bullet")
    for item in items:
        doc.add_paragraph(str(item), style=style) if style else doc.add_paragraph(f"– {item}")


def _set_footer_touchpoint(doc, tp: int) -> None:
    """Fusszeile der Vorlage: erste Zelle trägt 'Touchpoint 1' → aktueller TP."""
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            try:
                tables = footer.tables
            except Exception:
                continue
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if run.text.strip().lower().startswith("touchpoint"):
                                    run.text = f"Touchpoint {tp}"


def _title_block(doc, *, kicker: str, title: str, subtitle: str, meta: str) -> None:
    """Titelblock im Stil der Vorlage: Kicker (Title), Titel (Subtitle),
    Untertitel + Kurs-/Case-Zeile (Normal), 'St.Gallen, <Datum>'."""
    _para(doc, kicker, style=_style(doc, "Title"))
    _para(doc, title, style=_style(doc, "Subtitle"))
    _para(doc, subtitle)
    _para(doc, meta, grey=True, size=9.5)
    today = datetime.now(timezone.utc).date()
    _para(doc, f"St.Gallen, {_fmt_date_long(today)}", style=_style(doc, "Verfasser Ort Datum"))


# ---------------------------------------------------------------------------
# Produkt 1: KI-Briefing (ÜGL)
# ---------------------------------------------------------------------------

def _formal_table(doc, formal: dict, rubric: BriefingRubric) -> None:
    rows: list[tuple[str, str]] = []
    for key, label in (("baustein1", "Folie 2 (Baustein 1)"), ("baustein2", "Folie 3 (Baustein 2)")):
        chars = int(formal.get(f"{key}_chars", 0) or 0)
        limit = int(formal.get(f"{key}_max", rubric.max_chars(key)) or 0)
        status = "innerhalb der Grenze" if chars <= limit else f"über der Grenze (+{chars - limit})"
        rows.append((label, f"{chars:,} von {limit:,} Zeichen · {status}".replace(",", "'")))
    code = formal.get("code") or "nicht erkennbar"
    code_note = "gültig" if formal.get("code_valid") else "nicht im vorgegebenen Format"
    if formal.get("code") and not formal.get("code_matches_tp", True):
        code_note += ", Touchpoint im Code weicht ab"
    rows.append(("Code", f"{code} · {code_note}"))
    rows.append((
        "Dateiname",
        f"{formal.get('filename', '')} · "
        + ("entspricht dem Muster" if formal.get("filename_valid") else "weicht vom Muster ab"),
    ))
    fmt = str(formal.get("format", "")).upper()
    rows.append((
        "Format",
        fmt + (" · offizielle Vorlage erkannt" if formal.get("template_detected") else " · Vorlage nicht erkannt"),
    ))
    if formal.get("members_filled") is not None:
        rows.append(("Mitglieder", "ausgefüllt" if formal.get("members_filled") else "nicht ausgefüllt"))
    if formal.get("full_sentences_hint"):
        rows.append(("Satzform", str(formal["full_sentences_hint"])))

    table = doc.add_table(rows=0, cols=2)
    table.style = _style(doc, "Table Grid") or table.style
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for paragraph in cells[0].paragraphs + cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
    notes = [str(n) for n in formal.get("notes", []) if str(n).strip()]
    if notes:
        _para(doc, " ".join(notes), italic=True, size=9)


def _render_group(doc, record: dict, rubric: BriefingRubric) -> None:
    sg = record.get("sg")
    code = record.get("code") or record.get("filename", "")
    title = f"Stammgruppe SG{sg}" if sg else "Stammgruppe (nicht zugeordnet)"
    _heading(doc, f"{title} · {code}", 2)

    if record.get("status") == "extraction_failed":
        _para(
            doc,
            "Die Datei konnte nicht gelesen werden — bitte die Abgabe direkt öffnen. "
            + str(record.get("review_reason") or ""),
            italic=True,
        )
        return

    if record.get("needs_human_review"):
        reason = record.get("review_reason") or "Automatische Verdichtung mit Vorbehalt."
        _para(doc, f"Hinweis: {reason}", italic=True, size=9, grey=True)

    _heading(doc, "Formale Vorprüfung (gemeldet, nicht bewertet)", 3)
    _formal_table(doc, record.get("formal", {}) or {}, rubric)

    briefing = record.get("briefing", {}) or {}
    for b in rubric.bausteine:
        data = briefing.get(b.key, {}) or {}
        _heading(doc, f"Baustein {b.key[-1]} · {b.title} (Folie {b.slide}, Klausur {b.exam_ref})", 3)
        _para(doc, str(data.get("kernposition", "")), bold_label="Kernposition:")
        _para(doc, "", bold_label="Tragende Argumente:")
        _bullets(doc, list(data.get("tragende_argumente", []) or []), "Keine tragenden Argumente identifiziert.")
        _para(doc, "", bold_label="Dünne Stellen (Ansatz für Rückfragen):")
        _bullets(doc, list(data.get("duenne_stellen", []) or []), "Keine dünnen Stellen identifiziert.")
        _para(doc, str(data.get("einschaetzung", "")), bold_label="Einschätzung:")


def render_briefing_docx(
    records: list[dict],
    *,
    rubric: BriefingRubric,
    ueg: str,
    missing_groups: list[int] | None = None,
) -> bytes:
    """Rendert ein DOCX für eine Übungsgruppe (alle vorhandenen Stammgruppen)
    oder — bei genau einem Datensatz — für eine einzelne Stammgruppe."""
    doc = _new_document()
    _set_footer_touchpoint(doc, rubric.tp)
    schedule = BRIEFING_SCHEDULE.get(rubric.tp, {})
    single = len(records) == 1
    bausteine = " · ".join(b.title for b in rubric.bausteine)

    title = f"Touchpoint {rubric.tp} · Übungsgruppe {ueg or 'ohne Zuordnung'}"
    if single and records[0].get("sg"):
        title += f" · Stammgruppe SG{records[0]['sg']}"
    _title_block(
        doc,
        kicker="KI-Briefing",
        title=title,
        subtitle=bausteine,
        meta=(
            f"BWL A Assessment-Jahr HS26 · Running Case ON, Kapitel {rubric.case_chapter} · "
            f"Abgabe {_fmt_date(schedule.get('abgabe'))} · Termin {_fmt_date(schedule.get('termin'))} · "
            f"Klausurbezug {', '.join(rubric.exam_ref)} · Rubric {rubric.version} vom {rubric.date}"
        ),
    )

    _para(
        doc,
        "Nur für die Übungsgruppenleitung. Dieses Briefing verdichtet jede Abgabe entlang der "
        "Bausteine des Arbeitsauftrags: Kernposition, tragende Argumente, dünne Stellen. Es enthält "
        "keine Punkte, keine Stufen und keine Musterlösung — jede Wahl ist zulässig, beurteilt wird "
        "nur, ob die Begründung trägt. Die Wahl der Spannungslinie und der Rückfragen bleibt Ihre "
        "didaktische Entscheidung. Das Feedback an die Stammgruppen ist ein eigenes Dokument.",
        italic=True, size=9.5,
    )
    if missing_groups:
        p = _para(doc, "Keine Abgabe eingegangen: " + ", ".join(f"SG{n}" for n in missing_groups), size=9.5)
        for run in p.runs:
            run.bold = True

    ordered = sorted(records, key=lambda r: (r.get("sg") is None, int(r.get("sg") or 99), str(r.get("filename", ""))))
    for record in ordered:
        _render_group(doc, record, rubric)

    footer = _para(doc, "Automatisch erstellt durch ToAdapt · KI-Pipeline BWL A HS26", size=8, grey=True)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Produkt 2: KI-Feedback an die Stammgruppe (ein Dokument je Stammgruppe)
# ---------------------------------------------------------------------------

def render_feedback_docx(record: dict, *, rubric: BriefingRubric) -> bytes:
    """Rückmeldung an EINE Stammgruppe: je Baustein was trägt / was bleibt dünn /
    nächster Schritt, Abschluss Ausblick. Keine Punkte, keine Stufen, keine
    formale Vorprüfung, keine interne Einstufung — nur die Rückmeldung selbst."""
    doc = _new_document()
    _set_footer_touchpoint(doc, rubric.tp)
    schedule = BRIEFING_SCHEDULE.get(rubric.tp, {})
    sg = record.get("sg")
    code = record.get("code") or record.get("filename", "")

    title = f"Touchpoint {rubric.tp}"
    if sg:
        title += f" · Stammgruppe SG{sg}"
    _title_block(
        doc,
        kicker="Rückmeldung",
        title=title,
        subtitle=" · ".join(b.title for b in rubric.bausteine),
        meta=(
            f"BWL A Assessment-Jahr HS26 · Running Case ON, Kapitel {rubric.case_chapter} · Abgabe {code} · "
            f"Termin {_fmt_date(schedule.get('termin'))} · Klausurbezug {', '.join(rubric.exam_ref)}"
        ),
    )
    _para(
        doc,
        "Diese Rückmeldung bezieht sich ausschliesslich auf Ihr eigenes Ergebnis und folgt denselben "
        "Kriterien, die im Bewertungsraster der Klausur für die entsprechende Teilaufgabe gelten. Sie "
        "enthält keine Punkte und keine Musterlösung: Jede Wahl ist zulässig; es geht nur darum, wo "
        "Ihre Begründung trägt und wo sie dünn bleibt.",
        italic=True, size=9.5,
    )

    feedback = record.get("feedback", {}) or {}
    for b in rubric.bausteine:
        data = feedback.get(b.key, {}) or {}
        _heading(doc, f"Baustein {b.key[-1]} · {b.title} (Folie {b.slide}, Klausur {b.exam_ref})", 2)
        _para(doc, str(data.get("was_traegt", "")), bold_label="Was trägt:")
        _para(doc, str(data.get("was_bleibt_duenn", "")), bold_label="Was bleibt dünn:")
        _para(doc, str(data.get("naechster_schritt", "")), bold_label="Nächster Schritt:")

    _heading(doc, "Ausblick", 2)
    _para(doc, str(feedback.get("feed_forward", "")))

    footer = _para(
        doc,
        "Automatisch erstellt durch ToAdapt · KI-Pipeline BWL A HS26 · weitergegeben durch Ihre Übungsgruppenleitung",
        size=8, grey=True,
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
